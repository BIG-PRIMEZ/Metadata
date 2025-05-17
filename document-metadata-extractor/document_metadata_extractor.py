import os
import hashlib
import json
import sqlite3
import datetime
import tkinter as tk
from tkinter import filedialog, ttk, scrolledtext
import threading

# Document handler imports
import PyPDF2  # For PDFs
from docx import Document  # For Word documents
from PIL import Image  # For images (EXIF data)
from PIL.ExifTags import TAGS  # Add this import for EXIF tag names
import openpyxl  # For Excel files
from pptx import Presentation  # For PowerPoint files
import email  # For email files
import csv  # For CSV files

class DocumentMetadataExtractor:
    """Base class for document metadata extraction"""
    
    def __init__(self):
        self.supported_extensions = {
            '.pdf': self.extract_pdf_metadata,
            '.docx': self.extract_docx_metadata,
            '.xlsx': self.extract_xlsx_metadata,
            '.pptx': self.extract_pptx_metadata,
            '.jpg': self.extract_image_metadata,
            '.jpeg': self.extract_image_metadata,
            '.png': self.extract_image_metadata,
            '.eml': self.extract_email_metadata,
            '.csv': self.extract_csv_metadata,
        }
    
    def extract_metadata(self, file_path):
        """Extract metadata from a file based on its extension"""
        _, extension = os.path.splitext(file_path)
        extension = extension.lower()
        
        if extension in self.supported_extensions:
            return self.supported_extensions[extension](file_path)
        else:
            return {"error": f"Unsupported file type: {extension}"}
    
    def extract_pdf_metadata(self, file_path):
        """Extract metadata from PDF file"""
        metadata = {}
        try:
            with open(file_path, 'rb') as file:
                pdf = PyPDF2.PdfReader(file)
                if pdf.metadata:
                    for key, value in pdf.metadata.items():
                        # Remove the leading slash from PDF metadata keys
                        clean_key = key[1:] if key.startswith('/') else key
                        metadata[clean_key] = str(value)
                
                metadata['page_count'] = len(pdf.pages)
                # Get text from first page for content identification
                metadata['first_page_preview'] = pdf.pages[0].extract_text()[:200]
            return metadata
        except Exception as e:
            return {"error": f"Failed to extract PDF metadata: {str(e)}"}
    
    def extract_docx_metadata(self, file_path):
        """Extract metadata from Word document"""
        metadata = {}
        try:
            doc = Document(file_path)
            core_properties = doc.core_properties
            
            metadata['author'] = core_properties.author
            metadata['created'] = str(core_properties.created)
            metadata['last_modified_by'] = core_properties.last_modified_by
            metadata['modified'] = str(core_properties.modified)
            metadata['title'] = core_properties.title
            metadata['paragraph_count'] = len(doc.paragraphs)
            # Get text sample
            metadata['text_preview'] = doc.paragraphs[0].text[:200] if doc.paragraphs else ""
            
            return metadata
        except Exception as e:
            return {"error": f"Failed to extract Word metadata: {str(e)}"}
    
    def extract_xlsx_metadata(self, file_path):
        """Extract metadata from Excel file"""
        metadata = {}
        try:
            workbook = openpyxl.load_workbook(file_path, read_only=True)
            
            metadata['sheet_names'] = workbook.sheetnames
            metadata['sheet_count'] = len(workbook.sheetnames)
            metadata['creator'] = workbook.properties.creator
            metadata['created'] = str(workbook.properties.created)
            metadata['modified'] = str(workbook.properties.modified)
            metadata['last_modified_by'] = workbook.properties.lastModifiedBy
            
            return metadata
        except Exception as e:
            return {"error": f"Failed to extract Excel metadata: {str(e)}"}
    
    def extract_pptx_metadata(self, file_path):
        """Extract metadata from PowerPoint file"""
        metadata = {}
        try:
            presentation = Presentation(file_path)
            
            metadata['slide_count'] = len(presentation.slides)
            core_properties = presentation.core_properties
            metadata['author'] = core_properties.author
            metadata['created'] = str(core_properties.created)
            metadata['modified'] = str(core_properties.modified)
            metadata['title'] = core_properties.title
            
            return metadata
        except Exception as e:
            return {"error": f"Failed to extract PowerPoint metadata: {str(e)}"}
    
    def extract_image_metadata(self, file_path):
        """Extract EXIF metadata from image files"""
        metadata = {}
        try:
            with Image.open(file_path) as img:
                metadata['format'] = img.format
                metadata['size'] = f"{img.width}x{img.height}"
                metadata['mode'] = img.mode
                
                # Extract EXIF data if available
                exif_data = img._getexif()
                if exif_data:
                    for tag_id, value in exif_data.items():
                        tag = TAGS.get(tag_id, tag_id)
                        metadata[f"exif_{tag}"] = str(value)
            
            return metadata
        except Exception as e:
            return {"error": f"Failed to extract image metadata: {str(e)}"}
    
    def extract_email_metadata(self, file_path):
        """Extract metadata from email files"""
        metadata = {}
        try:
            with open(file_path, 'r') as file:
                msg = email.message_from_file(file)
                
                metadata['subject'] = msg.get('Subject', '')
                metadata['from'] = msg.get('From', '')
                metadata['to'] = msg.get('To', '')
                metadata['date'] = msg.get('Date', '')
                metadata['cc'] = msg.get('Cc', '')
                
                # Count attachments
                attachment_count = 0
                for part in msg.walk():
                    if part.get_content_disposition() == 'attachment':
                        attachment_count += 1
                
                metadata['attachment_count'] = attachment_count
            
            return metadata
        except Exception as e:
            return {"error": f"Failed to extract email metadata: {str(e)}"}
    
    def extract_csv_metadata(self, file_path):
        """Extract metadata from CSV files"""
        metadata = {}
        try:
            with open(file_path, 'r', newline='') as csvfile:
                csv_reader = csv.reader(csvfile)
                headers = next(csv_reader, None)
                
                # Count rows
                row_count = sum(1 for _ in csv_reader)
                
                metadata['column_count'] = len(headers) if headers else 0
                metadata['column_names'] = headers if headers else []
                metadata['row_count'] = row_count
            
            return metadata
        except Exception as e:
            return {"error": f"Failed to extract CSV metadata: {str(e)}"}


class MetadataHasher:
    """Hash the extracted metadata"""
    
    @staticmethod
    def hash_metadata(metadata):
        """Create a SHA-256 hash of the metadata"""
        # Convert metadata to a sorted JSON string to ensure consistent hashing
        metadata_str = json.dumps(metadata, sort_keys=True)
        return hashlib.sha256(metadata_str.encode()).hexdigest()


class DatabaseManager:
    """Manage the SQLite database operations"""
    
    def __init__(self, db_path="document_metadata.db"):
        self.db_path = db_path
        self.create_tables()
    
    def create_tables(self):
        """Create necessary tables if they don't exist"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        # Create table for document metadata
        cursor.execute('''
        CREATE TABLE IF NOT EXISTS document_metadata (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            file_path TEXT,
            file_name TEXT,
            metadata_hash TEXT,
            timestamp TEXT,
            metadata_json TEXT
        )
        ''')
        
        conn.commit()
        conn.close()
    
    def save_metadata(self, file_path, metadata, metadata_hash):
        """Save metadata and its hash to the database"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        timestamp = datetime.datetime.now().isoformat()
        file_name = os.path.basename(file_path)
        metadata_json = json.dumps(metadata)
        
        cursor.execute('''
        INSERT INTO document_metadata 
        (file_path, file_name, metadata_hash, timestamp, metadata_json)
        VALUES (?, ?, ?, ?, ?)
        ''', (file_path, file_name, metadata_hash, timestamp, metadata_json))
        
        conn.commit()
        conn.close()
        
        return cursor.lastrowid
    
    def get_all_metadata(self):
        """Retrieve all metadata records from the database"""
        conn = sqlite3.connect(self.db_path)
        conn.row_factory = sqlite3.Row
        cursor = conn.cursor()
        
        cursor.execute('SELECT * FROM document_metadata ORDER BY timestamp DESC')
        rows = cursor.fetchall()
        
        result = []
        for row in rows:
            result.append(dict(row))
        
        conn.close()
        return result


class MetadataExtractorGUI:
    """GUI for the metadata extraction system"""
    
    def __init__(self, master):
        self.master = master
        self.master.title("Document Metadata Extractor")
        self.master.geometry("900x700")
        
        self.extractor = DocumentMetadataExtractor()
        self.hasher = MetadataHasher()
        self.db_manager = DatabaseManager()
        
        self.current_file_path = None
        self.current_metadata = None
        
        self.create_widgets()
    
    def create_widgets(self):
        """Create and arrange GUI widgets"""
        # Frame for file selection
        file_frame = ttk.LabelFrame(self.master, text="Document Selection")
        file_frame.pack(fill="x", padx=10, pady=10)
        
        self.file_path_var = tk.StringVar()
        file_entry = ttk.Entry(file_frame, textvariable=self.file_path_var, width=70)
        file_entry.pack(side=tk.LEFT, padx=5, pady=5, fill="x", expand=True)
        
        browse_button = ttk.Button(file_frame, text="Browse", command=self.browse_file)
        browse_button.pack(side=tk.LEFT, padx=5, pady=5)
        
        extract_button = ttk.Button(file_frame, text="Extract Metadata", command=self.extract_metadata)
        extract_button.pack(side=tk.LEFT, padx=5, pady=5)
        
        # Frame for metadata display
        metadata_frame = ttk.LabelFrame(self.master, text="Extracted Metadata")
        metadata_frame.pack(fill="both", expand=True, padx=10, pady=10)
        
        self.metadata_text = scrolledtext.ScrolledText(metadata_frame, wrap=tk.WORD, width=80, height=20)
        self.metadata_text.pack(fill="both", expand=True, padx=5, pady=5)
        
        # Frame for hash display
        hash_frame = ttk.LabelFrame(self.master, text="Metadata Hash")
        hash_frame.pack(fill="x", padx=10, pady=10)
        
        self.hash_var = tk.StringVar()
        hash_entry = ttk.Entry(hash_frame, textvariable=self.hash_var, width=70, state="readonly")
        hash_entry.pack(fill="x", padx=5, pady=5)
        
        # Frame for database operations
        db_frame = ttk.LabelFrame(self.master, text="Database Operations")
        db_frame.pack(fill="x", padx=10, pady=10)
        
        save_button = ttk.Button(db_frame, text="Save to Database", command=self.save_to_database)
        save_button.pack(side=tk.LEFT, padx=5, pady=5)
        
        view_records_button = ttk.Button(db_frame, text="View All Records", command=self.view_all_records)
        view_records_button.pack(side=tk.LEFT, padx=5, pady=5)
        
        # Status bar
        self.status_var = tk.StringVar()
        self.status_var.set("Ready")
        status_bar = ttk.Label(self.master, textvariable=self.status_var, relief=tk.SUNKEN, anchor=tk.W)
        status_bar.pack(side=tk.BOTTOM, fill=tk.X)
    
    def browse_file(self):
        """Open file browser dialog to select a document"""
        file_path = filedialog.askopenfilename(
            title="Select a document",
            filetypes=(
                ("Document files", "*.pdf;*.docx;*.xlsx;*.pptx;*.jpg;*.jpeg;*.png;*.eml;*.csv"),
                ("PDF files", "*.pdf"),
                ("Word documents", "*.docx"),
                ("Excel files", "*.xlsx"),
                ("PowerPoint files", "*.pptx"),
                ("Image files", "*.jpg;*.jpeg;*.png"),
                ("Email files", "*.eml"),
                ("CSV files", "*.csv"),
                ("All files", "*.*")
            )
        )
        
        if file_path:
            self.file_path_var.set(file_path)
            self.current_file_path = file_path
            self.status_var.set(f"Selected file: {os.path.basename(file_path)}")
    
    def extract_metadata(self):
        """Extract metadata from the selected document"""
        if not self.current_file_path:
            self.status_var.set("Error: No file selected")
            return
        
        self.status_var.set("Extracting metadata...")
        
        # Run extraction in a separate thread to keep GUI responsive
        threading.Thread(target=self._extract_metadata_thread).start()
    
    def _extract_metadata_thread(self):
        """Thread function for metadata extraction"""
        try:
            # Extract metadata
            metadata = self.extractor.extract_metadata(self.current_file_path)
            
            # Calculate hash
            metadata_hash = self.hasher.hash_metadata(metadata)
            
            # Update GUI with results
            self.master.after(0, self._update_metadata_display, metadata, metadata_hash)
        except Exception as e:
            self.master.after(0, lambda: self.status_var.set(f"Error: {str(e)}"))
    
    def _update_metadata_display(self, metadata, metadata_hash):
        """Update GUI with extracted metadata and hash"""
        self.current_metadata = metadata
        
        # Display metadata in pretty format
        self.metadata_text.delete(1.0, tk.END)
        self.metadata_text.insert(tk.END, json.dumps(metadata, indent=4))
        
        # Display hash
        self.hash_var.set(metadata_hash)
        
        self.status_var.set("Metadata extraction complete")
    
    def save_to_database(self):
        """Save the extracted metadata to the database"""
        if not self.current_metadata:
            self.status_var.set("Error: No metadata to save")
            return
        
        try:
            metadata_hash = self.hash_var.get()
            record_id = self.db_manager.save_metadata(
                self.current_file_path, 
                self.current_metadata, 
                metadata_hash
            )
            
            self.status_var.set(f"Metadata saved to database with ID: {record_id}")
        except Exception as e:
            self.status_var.set(f"Error saving to database: {str(e)}")
    
    def view_all_records(self):
        """Open a new window to display all database records"""
        records = self.db_manager.get_all_metadata()
        
        # Create a new top-level window
        records_window = tk.Toplevel(self.master)
        records_window.title("Database Records")
        records_window.geometry("800x600")
        
        # Create a treeview to display the records
        columns = ("ID", "Filename", "Hash", "Timestamp")
        tree = ttk.Treeview(records_window, columns=columns, show="headings")
        
        # Set column headings
        for col in columns:
            tree.heading(col, text=col)
            tree.column(col, width=100)
        
        # Insert data
        for record in records:
            tree.insert("", "end", values=(
                record["id"],
                record["file_name"],
                record["metadata_hash"],
                record["timestamp"]
            ))
        
        # Add scrollbars
        vsb = ttk.Scrollbar(records_window, orient="vertical", command=tree.yview)
        hsb = ttk.Scrollbar(records_window, orient="horizontal", command=tree.xview)
        tree.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)
        
        # Pack everything
        tree.pack(fill="both", expand=True)
        vsb.pack(side="right", fill="y")
        hsb.pack(side="bottom", fill="x")
        
        # Add a detail view
        detail_frame = ttk.LabelFrame(records_window, text="Record Details")
        detail_frame.pack(fill="both", padx=10, pady=10)
        
        detail_text = scrolledtext.ScrolledText(detail_frame, height=10)
        detail_text.pack(fill="both", expand=True)
        
        # Show details when a record is selected
        def item_selected(event):
            selected_item = tree.selection()[0]
            item_id = tree.item(selected_item)["values"][0]
            
            for record in records:
                if record["id"] == item_id:
                    metadata = json.loads(record["metadata_json"])
                    detail_text.delete(1.0, tk.END)
                    detail_text.insert(tk.END, json.dumps(metadata, indent=4))
                    break
        
        tree.bind("<<TreeviewSelect>>", item_selected)



def run_application():
    """Run the document metadata extractor application"""
    root = tk.Tk()
    app = MetadataExtractorGUI(root)
    root.mainloop()

if __name__ == "__main__":
    run_application()